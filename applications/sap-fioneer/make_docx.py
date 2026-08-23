"""Build editable .docx versions of the CV and cover letter from their HTML sources."""
import re, sys
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x2E, 0x2E, 0x33)
GREY = RGBColor(0x3A, 0x3A, 0x40)
SERIF = "Georgia"
SANS = "Arial"


class Node:
    def __init__(self, tag, attrs=None):
        self.tag, self.attrs, self.kids, self.parent = tag, dict(attrs or {}), [], None

    def add(self, n):
        if isinstance(n, Node):
            n.parent = self
        self.kids.append(n)

    def cls(self):
        return self.attrs.get("class", "").split()

    def text(self):
        out = []
        for k in self.kids:
            out.append(k if isinstance(k, str) else k.text())
        return "".join(out)


class Tree(HTMLParser):
    VOID = {"br", "meta", "link", "img", "hr"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = Node("root")
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        if tag in self.VOID:
            return
        n = Node(tag, attrs)
        self.stack[-1].add(n)
        self.stack.append(n)

    def handle_endtag(self, tag):
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                return

    def handle_data(self, data):
        self.stack[-1].add(data)


def clean(s):
    return re.sub(r"\s+", " ", s.replace(" ", " ")).strip()


def add_hyperlink(par, url, runs):
    """Attach an external-relationship hyperlink carrying already-built runs."""
    rid = par.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    for r in runs:
        par._p.remove(r._r)
        link.append(r._r)
    par._p.append(link)


def underline(run):
    run.font.underline = True
    return run


def emit_runs(par, node, bold=False, italic=False, size=None, color=None, font=SERIF):
    """Walk inline content, preserving <b>/<i>/<a> and .k key spans."""
    for k in node.kids:
        if isinstance(k, str):
            t = re.sub(r"\s+", " ", k.replace(" ", " "))
            if not t.strip() and not par.runs:
                continue
            r = par.add_run(t)
            r.bold, r.italic, r.font.name = bold, italic, font
            if size:
                r.font.size = size
            if color:
                r.font.color.rgb = color
        else:
            kc = k.cls()
            b = bold or k.tag in ("b", "strong") or "k" in kc
            i = italic or k.tag in ("i", "em")
            c = NAVY if (k.tag in ("b", "strong") or "k" in kc) and color is None else color
            f = SANS if "k" in kc else font
            sz = Pt(8.3) if "k" in kc else size
            if k.tag == "a" and k.attrs.get("href"):
                before = len(par.runs)
                emit_runs(par, k, b, i, sz, c, f)
                new = par.runs[before:]
                for r in new:
                    underline(r)
                if new:
                    add_hyperlink(par, k.attrs["href"], new)
            else:
                emit_runs(par, k, b, i, sz, c, f)


def spacing(par, before=0, after=2, line=1.16):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def bottom_border(par, sz=6, color="000000"):
    p = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), color)
    pbdr.append(bot)
    p.append(pbdr)


def letter_space(run, twentieths):
    """w:spacing takes twentieths of a point."""
    rPr = run._r.get_or_add_rPr()
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(twentieths))
    rPr.append(el)


def top_border(par, sz=6, color="000000"):
    p = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(sz))
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), color)
    pbdr.append(top)
    p.append(pbdr)


def right_tab(par, doc):
    sec = doc.sections[0]
    width = sec.page_width - sec.left_margin - sec.right_margin
    par.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)


def find(node, pred, out=None):
    out = out if out is not None else []
    for k in node.kids:
        if isinstance(k, Node):
            if pred(k):
                out.append(k)
            find(k, pred, out)
    return out


def setup(doc, margins):
    st = doc.styles["Normal"]
    st.font.name = SERIF
    st.font.size = Pt(9.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
    sec = doc.sections[0]
    sec.top_margin, sec.right_margin, sec.bottom_margin, sec.left_margin = margins
    for s in ("List Bullet",):
        try:
            doc.styles[s].font.name = SERIF
        except KeyError:
            pass


def header_block(doc, tree, body_size):
    h1 = find(tree.root, lambda n: n.tag == "h1")[0]
    p = doc.add_paragraph()
    spacing(p, 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(clean(h1.text()).upper())
    r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(19), NAVY
    letter_space(r, 130)

    tag = find(tree.root, lambda n: "tagline" in n.cls())[0]
    p = doc.add_paragraph()
    spacing(p, 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(clean(tag.text()).upper())
    r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(8.5), BLUE
    letter_space(r, 70)
    top_border(p, sz=12)

    con = find(tree.root, lambda n: "contact" in n.cls())[0]
    p = doc.add_paragraph()
    spacing(p, 0, 6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emit_runs(p, con, size=Pt(8.5), color=GREY)
    top_border(p, sz=4)


def build_cv(src, dst):
    tree = Tree()
    tree.feed(open(src, encoding="utf-8").read())
    doc = Document()
    setup(doc, (Cm(1.1), Cm(1.3), Cm(1.0), Cm(1.3)))
    header_block(doc, tree, Pt(9.5))

    for sec in find(tree.root, lambda n: n.tag == "section"):
        h2 = find(sec, lambda n: n.tag == "h2")
        if h2:
            p = doc.add_paragraph()
            spacing(p, 7, 3)
            r = p.add_run(clean(h2[0].text()).upper())
            r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(8.6), NAVY
            r.font.all_caps = True
            letter_space(r, 95)
            bottom_border(p, sz=8)

        for blk in sec.kids:
            if not isinstance(blk, Node):
                continue
            c = blk.cls()
            if blk.tag == "h2":
                continue
            if blk.tag == "p":
                p = doc.add_paragraph()
                spacing(p, 0, 3)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                emit_runs(p, blk, size=Pt(9.4))
            elif "comp" in c or "skills" in c:
                for d in blk.kids:
                    if isinstance(d, Node) and d.tag == "div":
                        p = doc.add_paragraph()
                        spacing(p, 0, 2)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        emit_runs(p, d, size=Pt(8.9))
            elif "role" in c:
                head = find(blk, lambda n: "role-head" in n.cls())[0]
                title = find(head, lambda n: "role-title" in n.cls())[0]
                date = find(head, lambda n: "role-date" in n.cls())[0]
                p = doc.add_paragraph()
                spacing(p, 5, 1)
                right_tab(p, doc)
                r = p.add_run(clean(title.text()))
                r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(9.8), NAVY
                r = p.add_run("\t" + clean(date.text()).upper())
                r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(8.2), GREY
                letter_space(r, 25)
                sub = find(blk, lambda n: "role-sub" in n.cls())
                if sub:
                    p = doc.add_paragraph()
                    spacing(p, 0, 2)
                    r = p.add_run(clean(sub[0].text()))
                    r.italic, r.font.size, r.font.color.rgb = True, Pt(8.8), GREY
                for li in find(blk, lambda n: n.tag == "li"):
                    p = doc.add_paragraph(style="List Bullet")
                    spacing(p, 0, 2)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.left_indent = Cm(0.5)
                    emit_runs(p, li, size=Pt(9.3))
            elif "leave" in c:
                p = doc.add_paragraph()
                spacing(p, 3, 2)
                right_tab(p, doc)
                spans = [k for k in blk.kids if isinstance(k, Node)]
                r = p.add_run(clean(spans[0].text()))
                r.italic, r.font.size, r.font.color.rgb = True, Pt(8.8), GREY
                r = p.add_run("\t" + clean(spans[1].text()))
                r.italic, r.font.size, r.font.color.rgb = True, Pt(8.8), GREY
            elif "proj" in c:
                head = find(blk, lambda n: "proj-head" in n.cls())[0]
                name = find(head, lambda n: "proj-name" in n.cls())[0]
                date = find(head, lambda n: "proj-date" in n.cls())[0]
                links = find(head, lambda n: "links" in n.cls())
                p = doc.add_paragraph()
                spacing(p, 3, 0)
                right_tab(p, doc)
                r = p.add_run(clean(name.text()))
                r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(9), NAVY
                for a in (find(links[0], lambda n: n.tag == "a") if links else []):
                    p.add_run("  ").font.size = Pt(8)
                    r = p.add_run(clean(a.text()))
                    r.bold, r.font.name, r.font.size = True, SANS, Pt(7.8)
                    r.font.color.rgb, r.font.underline = NAVY, True
                    add_hyperlink(p, a.attrs["href"], [r])
                r = p.add_run("\t" + clean(date.text()))
                r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(8.3), GREY
                for pp in find(blk, lambda n: n.tag == "p"):
                    q = doc.add_paragraph()
                    spacing(q, 0, 2)
                    q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    emit_runs(q, pp, size=Pt(9))
            elif "edu-row" in c:
                spans = [k for k in blk.kids if isinstance(k, Node)]
                p = doc.add_paragraph()
                spacing(p, 0, 2)
                right_tab(p, doc)
                emit_runs(p, spans[0], size=Pt(9.3))
                r = p.add_run("\t" + clean(spans[1].text()))
                r.font.size, r.font.color.rgb = Pt(8.5), GREY
    doc.save(dst)


def build_letter(src, dst):
    tree = Tree()
    tree.feed(open(src, encoding="utf-8").read())
    doc = Document()
    setup(doc, (Cm(1.6), Cm(1.8), Cm(1.4), Cm(1.8)))
    doc.styles["Normal"].font.size = Pt(10.3)
    header_block(doc, tree, Pt(10.3))

    meta = find(tree.root, lambda n: "meta" in n.cls())[0]
    for d in meta.kids:
        if isinstance(d, Node):
            p = doc.add_paragraph()
            spacing(p, 0, 2)
            emit_runs(p, d, size=Pt(9.8), color=GREY)

    subj = find(tree.root, lambda n: "subject" in n.cls())[0]
    p = doc.add_paragraph()
    spacing(p, 8, 8)
    r = p.add_run(clean(subj.text()))
    r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(10.2), NAVY

    body = tree.root
    for blk in find(body, lambda n: n.tag in ("p", "li")):
        if any(x in ("meta", "subject", "contact", "tagline", "sub") for x in blk.cls()):
            continue
        if blk.parent is not None and "sign" in blk.parent.cls():
            continue
        if blk.tag == "li":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6)
            spacing(p, 0, 4, 1.25)
        else:
            p = doc.add_paragraph()
            spacing(p, 0, 8, 1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        emit_runs(p, blk, size=Pt(10.3))

    sign = find(tree.root, lambda n: "sign" in n.cls())[0]
    p = doc.add_paragraph()
    spacing(p, 10, 2)
    p.add_run("Yours sincerely,").font.size = Pt(10.3)
    p = doc.add_paragraph()
    spacing(p, 14, 1)
    r = p.add_run("Mehreen Himani")
    r.bold, r.font.name, r.font.size, r.font.color.rgb = True, SANS, Pt(10.5), NAVY
    sub = find(sign, lambda n: "sub" in n.cls())[0]
    p = doc.add_paragraph()
    spacing(p, 0, 0)
    r = p.add_run(clean(sub.text()))
    r.font.size, r.font.color.rgb = Pt(9), GREY
    doc.save(dst)


if __name__ == "__main__":
    a = sys.argv[1:]
    cv_src = a[0] if len(a) > 0 else "cv.html"
    cv_out = a[1] if len(a) > 1 else "Mehreen_Himani_CV_SAP_Fioneer_Senior_Solution_Manager.docx"
    cl_src = a[2] if len(a) > 2 else "cover_letter.html"
    cl_out = a[3] if len(a) > 3 else "Mehreen_Himani_Cover_Letter_SAP_Fioneer.docx"
    build_cv(cv_src, cv_out)
    build_letter(cl_src, cl_out)
    print("docx written:", cv_out, cl_out)
